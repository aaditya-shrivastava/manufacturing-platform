from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from datetime import datetime

def export_to_docx(markdown_content: str, query: str, output_path: str):
    """Convert markdown report to a professional .docx file."""
    doc = Document()

    # ── Title ──────────────────────────────────────────────
    title = doc.add_heading("Manufacturing Procurement Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Metadata ───────────────────────────────────────────
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Query: {query}  |  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    # ── Parse and write markdown lines ────────────────────
    lines = markdown_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)

        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)

        elif line.startswith("# "):
            pass  # skip — already added as title

        elif line.startswith("| ") and "|" in line:
            # Collect all table rows
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                if not re.match(r"^\|[\s\-|]+\|$", row):  # skip separator rows
                    table_lines.append(row)
                i += 1

            if table_lines:
                cols = [c.strip() for c in table_lines[0].split("|") if c.strip()]
                table = doc.add_table(rows=1, cols=len(cols))
                table.style = "Light Grid Accent 1"

                # Header row
                for j, col in enumerate(cols):
                    cell = table.rows[0].cells[j]
                    cell.text = col
                    run = cell.paragraphs[0].runs[0]
                    run.bold = True

                # Data rows
                for row_line in table_lines[1:]:
                    cells = [c.strip() for c in row_line.split("|") if c.strip()]
                    row = table.add_row()
                    for j, val in enumerate(cells[:len(cols)]):
                        row.cells[j].text = val

                doc.add_paragraph()
            continue

        elif line.startswith("* ") or line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")

        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")

        elif line == "":
            pass  # skip blank lines

        else:
            # Clean bold markdown (**text**) before adding
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            doc.add_paragraph(clean)

        i += 1

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"📄 Word report saved to: {output_path}")